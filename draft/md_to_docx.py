"""把 draft/final_report.md 转换为 docx 文档。
依赖：python-docx >= 1.0
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "draft" / "final_report.md"
DOCX_PATH = ROOT / "结题报告_第4组.docx"


# ---------- 工具函数 ----------

def set_cell_bg(cell, hex_color: str) -> None:
    """给表格单元格设置底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cn_font(run, size_pt: int | None = None, bold: bool = False) -> None:
    """设置中文字体（东亚字体）。"""
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold


# ---------- Markdown 解析 ----------

INLINE_RE = re.compile(r"\*\*(.+?)\*\*|`([^`]+)`")


def render_inline(paragraph, text: str, size_pt: int = 11, bold_default: bool = False) -> None:
    """把含 **加粗** 与 `code` 的 Markdown 行内语法渲染到段落。"""
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_cn_font(run, size_pt, bold_default)
        if match.group(1) is not None:  # **bold**
            run = paragraph.add_run(match.group(1))
            set_cn_font(run, size_pt, True)
        else:  # `code`
            run = paragraph.add_run(match.group(2))
            run.font.name = "Consolas"
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_cn_font(run, size_pt, bold_default)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-+:?", c) for c in cells) and len(cells) >= 2


def render_table(doc: Document, rows: list[list[str]], header_idx: int = 0) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            render_inline(para, cell_text, size_pt=10, bold_default=(r_idx == header_idx))
            if r_idx == header_idx:
                set_cell_bg(cell, "DCE6F1")

    # 给整张表格字体兜底
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_cn_font(run, 10, run.bold)


def render_cover_table(doc: Document, rows: list[list[str]]) -> None:
    """封面内的字段表格：无表头底色、左列窄右列宽、整体细边框、字段列右对齐。"""
    if not rows:
        return
    # 直接用第一行作为字段行（cover 内的 markdown table 第一行其实是第一个字段，不是表头）
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 关闭 Word 默认网格线，再单独画细边框
    table.autofit = False
    widths = [Cm(3.5), Cm(11.0)]
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.width = widths[c_idx] if c_idx < len(widths) else Cm(4.0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            render_inline(para, cell_text, size_pt=12, bold_default=False)
            # 字段列加灰色细底
            if c_idx == 0:
                set_cell_bg(cell, "F2F2F2")

    # 加灰色细边框
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "6")
                b.set(qn("w:color"), "808080")
                tc_borders.append(b)
            tc_pr.append(tc_borders)

    # 字体兜底
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_cn_font(run, 12, False)


def parse_md(md_text: str) -> list[tuple[str, object]]:
    """把 Markdown 文本解析为 (kind, payload) 流。
    kind 可取值：heading1~4, paragraph, table, code, hr, blank, cover_*, right_para, cover_table.
    """
    lines = md_text.splitlines()
    blocks: list[tuple[str, object]] = []
    i = 0
    n = len(lines)
    in_cover = False
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 封面标记
        if stripped == "<!-- COVER_START -->":
            blocks.append(("cover_start", ""))
            in_cover = True
            i += 1
            continue
        if stripped == "<!-- COVER_END -->":
            blocks.append(("cover_end", ""))
            in_cover = False
            i += 1
            continue

        # 空行
        if not stripped:
            blocks.append(("blank", ""))
            i += 1
            continue

        # 分隔线
        if re.fullmatch(r"-{3,}", stripped):
            blocks.append(("hr", ""))
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            blocks.append((f"heading{level}", m.group(2).strip()))
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # 跳过结尾 ```
            blocks.append(("code", "\n".join(buf)))
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            header = parse_table_row(line)
            i += 2
            rows = [header]
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            blocks.append(("cover_table" if in_cover else "table", rows))
            continue

        # 右对齐 paragraph（仅在封面内有效）
        if in_cover and stripped.startswith("[RIGHT]"):
            blocks.append(("right_para", stripped[len("[RIGHT]"):].lstrip()))
            i += 1
            continue

        # 无序列表
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]).rstrip())
                i += 1
            blocks.append(("bullet", items))
            continue

        # 普通段落：合并连续非空行
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not (
            lines[i].lstrip().startswith(("#", "|", "- ", "* ", "```"))
            or re.fullmatch(r"-{3,}", lines[i].strip())
        ):
            buf.append(lines[i].strip())
            i += 1
        blocks.append(("paragraph", " ".join(buf)))

    return blocks


# ---------- 渲染到 docx ----------

def render(doc: Document, blocks: list[tuple[str, object]]) -> None:
    in_cover = False
    for kind, payload in blocks:
        if kind == "cover_start":
            in_cover = True
            # 封面顶部留白
            for _ in range(3):
                doc.add_paragraph()
            continue
        if kind == "cover_end":
            in_cover = False
            doc.add_page_break()
            continue
        if kind == "right_para":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            render_inline(p, str(payload), size_pt=12)
            continue
        if kind == "blank":
            if in_cover:
                # 封面内空行作为段间距
                doc.add_paragraph()
            continue
        if kind == "hr":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("―" * 30)
            continue
        if kind.startswith("heading"):
            level = int(kind[-1])
            if in_cover:
                # 封面内层级：h1=28pt（大学名、底部日期），h2=32pt（论文类型主标题），h3=16pt（学期）
                cover_sizes = {1: 28, 2: 32, 3: 16, 4: 14}
                text = str(payload)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_cn_font(run, cover_sizes.get(level, 14), True)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                if level <= 2:
                    doc.add_paragraph()
            else:
                sizes = {1: 18, 2: 15, 3: 13, 4: 12}
                text = str(payload)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_cn_font(run, sizes.get(level, 11), True)
                if level == 1:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                elif level == 2:
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                # 标题前后留白
                if level <= 2:
                    doc.add_paragraph()
            continue
        if kind == "paragraph":
            if in_cover:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(2.0)
                p.paragraph_format.space_after = Pt(4)
                render_inline(p, str(payload), size_pt=14)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                render_inline(p, str(payload), size_pt=11)
            continue
        if kind == "cover_table":
            render_cover_table(doc, payload)  # type: ignore[arg-type]
            doc.add_paragraph()
            continue
        if kind == "table":
            render_table(doc, payload, header_idx=0)  # type: ignore[arg-type]
            doc.add_paragraph()
            continue
        if kind == "code":
            for line in str(payload).splitlines():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(line if line else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue
        if kind == "bullet":
            for item in payload:  # type: ignore[union-attr]
                p = doc.add_paragraph(style="List Bullet")
                render_inline(p, item, size_pt=11)
            continue


def build_document(md_path: Path, out_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # 全局样式
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    blocks = parse_md(md_text)
    render(doc, blocks)
    doc.save(out_path)


if __name__ == "__main__":
    build_document(MD_PATH, DOCX_PATH)
    print(f"已生成：{DOCX_PATH}")